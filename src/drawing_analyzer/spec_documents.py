"""Extraction + size-budgeting for uploaded project specification documents.

Accepted formats: .pdf (pypdf), .docx (python-docx), .txt/.md (read directly).
No PyMuPDF import (I-5) — this module only ever produces plain text, never
touches page geometry or rendering, so it stays outside the AGPL-isolation
boundary that confines PyMuPDF to :mod:`drawing_analyzer.render` and
:mod:`drawing_analyzer.annotate`.

Every extraction failure is captured on the returned :class:`SpecDocument`
(I-3: additive, non-fatal) — this module never raises for a bad input file.

This is a *different, unrelated* concept from "Project Context" as used
elsewhere in this repo (:data:`core.tokenizer.PROJECT_CONTEXT_MAX_TOKENS`,
and docstrings in ``pipeline.py``/``export.py``) — that term names an external,
sibling spec-review tool this repo's ``combined_text`` output is manually
pasted into. This module's output is the operator's uploaded project
*specifications* — ground-truth reference material for THIS run, folded only
into the digest system prompt (see ``digest.py``'s ``_SPECS_ADDENDUM_TEMPLATE``).
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

SPEC_FILE_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md"})

# Per-file cap: one huge spec document can't crowd out the others when several
# are uploaded together. Whole-block cap: bounds total system-prompt growth
# (and therefore first-sheet cache-write cost + attention dilution) regardless
# of how many files were attached. See CLAUDE.md-adjacent design notes in the
# project plan for the cost/coverage rationale behind these numbers.
SPEC_FILE_CHAR_BUDGET = 40_000
SPEC_TOTAL_CHAR_BUDGET = 120_000


@dataclass
class SpecDocument:
    """One uploaded spec file's extraction result. Never raises — a bad file
    (unsupported type, corrupt, empty/unreadable) is captured in ``error``."""

    path: Path
    display_name: str
    text: str = ""
    error: str | None = None

    @property
    def ok(self) -> bool:
        return self.error is None and bool(self.text.strip())


def _extract_pdf_text(path: Path) -> str:
    import pypdf  # lazy: keep this module importable without the optional dep

    reader = pypdf.PdfReader(str(path))
    if getattr(reader, "is_encrypted", False):
        # Try an empty-password unlock (common for "owner-password-only"
        # protected PDFs); a genuinely user-protected file still raises below.
        reader.decrypt("")
    return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _extract_docx_text(path: Path) -> str:
    import docx  # python-docx; lazy import

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:  # spec tables (schedules, submittal matrices) matter
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace").strip()


_EXTRACTORS = {
    ".pdf": _extract_pdf_text,
    ".docx": _extract_docx_text,
    ".txt": _extract_plain_text,
    ".md": _extract_plain_text,
}


def extract_spec_text(path: Path) -> SpecDocument:
    """Extract one spec file's text. NEVER raises (I-3) — a bad file degrades
    to a populated ``SpecDocument.error``; the caller decides what to do next."""
    path = Path(path)
    suffix = path.suffix.lower()
    extractor = _EXTRACTORS.get(suffix)
    if extractor is None:
        return SpecDocument(
            path=path,
            display_name=path.name,
            error=(
                f"unsupported file type {suffix or '(none)'} "
                f"(accepted: {', '.join(sorted(SPEC_FILE_EXTENSIONS))})"
            ),
        )
    try:
        text = extractor(path)
    except Exception as exc:  # noqa: BLE001 - one bad spec file must never sink the upload
        return SpecDocument(path=path, display_name=path.name, error=str(exc)[:300])
    if not text.strip():
        return SpecDocument(
            path=path,
            display_name=path.name,
            error="no extractable text (empty, scanned/image-only, or unreadable)",
        )
    return SpecDocument(path=path, display_name=path.name, text=text)


def extract_spec_documents(paths: list[Path]) -> list[SpecDocument]:
    """Extract every path; per-file failures never abort the batch (I-3)."""
    return [extract_spec_text(p) for p in paths]


@dataclass
class SpecBudget:
    """Mirrors :class:`cross_qc._Budget` (``cross_qc.py``): omissions are
    *counted and surfaced*, never silently dropped."""

    total_chars: int = 0
    included_chars: int = 0
    omitted_chars: int = 0
    omitted_files: list[str] = field(default_factory=list)
    budget_chars: int = SPEC_TOTAL_CHAR_BUDGET

    @property
    def degraded(self) -> bool:
        return self.omitted_chars > 0


def _budgeted(text: str, budget_chars: int) -> tuple[str, int]:
    """Truncate ``text`` to ``budget_chars``; return ``(kept_text, omitted_count)``.

    The returned text NEVER exceeds ``budget_chars`` — the truncation marker
    is carved out of the budget, not appended on top of it. This matters
    because ``build_specs_text``'s whole-block truncation and
    ``enforce_specs_budget``'s pipeline-side backstop apply the *same*
    default budget one after another (the backstop protects a direct library
    caller who bypassed ``build_specs_text``); if a marker could push the
    result past the budget, the second pass could truncate mid-marker,
    sending a garbled ``...[TRUNCATED 452`` fragment to the model.
    """
    if len(text) <= budget_chars:
        return text, 0
    omitted = len(text) - budget_chars
    marker = f"\n[TRUNCATED {omitted} chars]"
    if len(marker) >= budget_chars:
        return text[:budget_chars], omitted  # degenerate: no room for a marker at all
    return text[: budget_chars - len(marker)] + marker, omitted


def build_specs_text(documents: list[SpecDocument]) -> tuple[str, SpecBudget]:
    """Concatenate every successfully-extracted document under a per-file
    header, applying the per-file budget first (so one huge file can't starve
    the others) and then the whole-block budget (a hard ceiling regardless of
    file count). Failed/empty documents are skipped here — their errors are
    the caller's responsibility to surface (the GUI logs them per-file)."""
    budget = SpecBudget()
    parts: list[str] = []
    for doc in documents:
        if not doc.ok:
            continue
        budget.total_chars += len(doc.text)
        kept, omitted = _budgeted(doc.text, SPEC_FILE_CHAR_BUDGET)
        if omitted:
            budget.omitted_chars += omitted
            budget.omitted_files.append(doc.display_name)
        parts.append(f"===== {doc.display_name} =====\n{kept}")
    combined = "\n\n".join(parts).strip()
    combined, extra_omitted = _budgeted(combined, SPEC_TOTAL_CHAR_BUDGET)
    if extra_omitted:
        budget.omitted_chars += extra_omitted
    # The true length of what's actually returned — not an accumulated,
    # per-file running total, which would miss the "===== name =====\n"
    # header/separator overhead and any whole-block truncation above.
    budget.included_chars = len(combined)
    return combined, budget


def enforce_specs_budget(
    specs_text: "object | None", *, budget_chars: int = SPEC_TOTAL_CHAR_BUDGET
) -> tuple[str, SpecBudget]:
    """Defensive backstop applied inside ``extract_drawing_context`` itself
    (I-3): a direct library caller who never went through the GUI's
    :func:`build_specs_text` still gets the total-chars ceiling enforced, so
    an arbitrarily large string handed straight to the pipeline can't blow the
    per-sheet system-prompt cost/attention budget.

    Coerces via ``str()`` (mirroring :func:`digest.normalize_specs_text`)
    rather than assuming ``specs_text`` is already a string — never raises on
    a caller mistake (e.g. passing a list of doc texts instead of joining
    them first), consistent with I-3."""
    text = "" if specs_text is None else str(specs_text).strip()
    if not text:
        return "", SpecBudget(budget_chars=budget_chars)
    kept, omitted = _budgeted(text, budget_chars)
    budget = SpecBudget(
        total_chars=len(text),
        included_chars=len(kept),
        omitted_chars=omitted,
        budget_chars=budget_chars,
    )
    return kept, budget
